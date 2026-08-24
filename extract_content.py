#!/usr/bin/env python3
"""Extract full document content with structure for analysis."""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX = "OCP_eGuide_Rapport_Stage_Final_CORRECTED.docx"

doc = Document(DOCX)

print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Sections: {len(doc.sections)}")

# Extract all paragraphs with styles
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else "None"
    text = para.text.strip()
    if text:
        print(f"\n[{i:03d}] <{style}> {text[:200]}")
    elif style.startswith("Heading"):
        print(f"\n[{i:03d}] <{style}> (empty heading)")

# Extract tables
print("\n\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti+1}: {len(table.rows)} rows x {len(table.columns)} cols")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:40] for cell in row.cells]
        print(f"  Row {ri}: {' | '.join(cells)}")
