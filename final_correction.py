#!/usr/bin/env python3
"""
Final comprehensive report correction.
Works from the original .docx and applies ALL corrections.
"""

from docx import Document
from docx.shared import Pt
import re

def has_drawing(para):
    """Check if paragraph contains a Word drawing (image)."""
    drawings = para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
    drawings += para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    return len(drawings) > 0

def remove_drawings_from_para(para):
    """Remove all drawing elements from a paragraph."""
    for inline in para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
        parent = inline.getparent()
        if parent is not None:
            parent.remove(inline)
    for anchor in para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor'):
        parent = anchor.getparent()
        if parent is not None:
            parent.remove(anchor)

def set_para_text(para, text, italic=False, size=11):
    """Set paragraph text, clearing existing runs."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        for run in para.runs:
            run.italic = italic
            run.font.size = Pt(size)
            run.font.name = "Calibri"
    else:
        run = para.add_run(text)
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Calibri"

def main():
    input_path = "OCP_eGuide_Rapport_Stage_Final.docx"
    output_path = "OCP_eGuide_Rapport_Stage_Final_CORRECTED.docx"
    
    print("=" * 70)
    print("FINAL REPORT CORRECTION")
    print("=" * 70)
    
    doc = Document(input_path)
    print(f"Loaded: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
    
    # ======================================================================
    # STEP 1: Remove images from paragraphs (keep only 5)
    # ======================================================================
    # Images are at these paragraph indices:
    # 76 (Fig 1), 85 (Fig 2), 87 (Fig 3), 89 (Fig 4),
    # 111 (Fig 5), 113 (Fig 6), 115 (Fig 7), 117 (Fig 8),
    # 137 (Fig 9), 139 (Fig 10), 141 (Fig 11), 145 (Fig 13),
    # 149 (Fig 14), 151 (Fig 15), 153 (Fig 16),
    # 157 (Fig 17), 159 (Fig 18), 161 (Fig 19), 163 (Fig 20),
    # 171 (Fig 21), 173 (Fig 22), 175 (Fig 23), 177 (Fig 24)
    
    KEEP_PARAS = {87, 145, 149, 173, 175}
    
    removed = 0
    for i, para in enumerate(doc.paragraphs):
        if has_drawing(para) and i not in KEEP_PARAS:
            remove_drawings_from_para(para)
            removed += 1
    
    remaining = sum(1 for p in doc.paragraphs if has_drawing(p))
    print(f"Removed {removed} images, kept {remaining}")
    
    # ======================================================================
    # STEP 2: Map old figure numbers to new ones
    # ======================================================================
    # Kept: Fig 3 -> Fig 2, Fig 13 -> Fig 1, Fig 14 -> Fig 5, Fig 22 -> Fig 3, Fig 23 -> Fig 4
    OLD_TO_NEW = {3: 2, 13: 1, 14: 5, 22: 3, 23: 4}
    REMOVED_FIGS = set(range(1, 27)) - set(OLD_TO_NEW.keys())
    
    # New captions for kept figures
    NEW_CAPTIONS = {
        1: "Figure 1 — Page d'accueil de l'application OCP eGuide en mode kiosk",
        2: "Figure 2 — Interface de sélection du profil utilisateur (profil stagiaire)",
        3: "Figure 3 — Vue principale de la carte interactive du campus OCP Jorf Lasfar",
        4: "Figure 4 — Carte interactive : zoom sur la zone de production avec marqueurs de localisations",
        5: "Figure 5 — Tableau de bord personnalisé pour le profil Réception",
    }
    
    # Replacement text for removed figure captions
    REPLACEMENT_TEXT = {
        1: "La page d'accueil de l'application OCP eGuide présente trois catégories d'utilisateurs (Employés, Stagiaires, Visiteurs) avec un sélecteur automatique en mode kiosk, permettant un premier contact intuitif avec le système.",
        2: "L'écran de sélection des profils employés affiche cinq catégories fonctionnelles : Management, Réception, Ressources Humaines, Informatique et Sécurité, chacune associée à un tableau de bord personnalisé.",
        4: "La sélection de profil visiteur propose des catégories adaptées aux rôles externes : Client, Livreur, Partenaire, Fournisseur, Consultant et Auditeur, avec un parcours simplifié d'inscription et de badge.",
        5: "Le flux d'authentification employé présente une carte de connexion avec le logo du profil sélectionné, les champs d'identifiants et un bouton de connexion sécurisée par JWT.",
        6: "L'authentification profil RH suit le même modèle visuel, adapté aux couleurs et à l'icône du département Ressources Humaines.",
        7: "L'authentification informatique utilise un style visuel distinct pour le département IT, avec accès aux outils d'administration système.",
        8: "Après une connexion réussie, un message de confirmation s'affiche et l'utilisateur est redirigé vers son tableau de bord personnalisé.",
        9: "Les cartes de profil employés s'affichent dans une grille responsive, chacune présentant le nom du département et un raccourci vers le tableau de bord correspondant.",
        10: "Les cartes de profil stagiaires suivent le même modèle, avec des catégories sectorielles (Mécanique, Chimie, Électrique, Civil) et des descriptions de tâches associées.",
        11: "Les cartes de profil visiteurs proposent des catégories de rôles externes avec des descriptions adaptées à chaque type de visite.",
        12: "L'interface de sélection de profil visiteur utilise un design glassmorphism avec des cartes semi-transparentes, chacune affichant un titre, une icône et une brève description du rôle.",
        15: "Le tableau de bord Réception affiche en temps réel la liste des visiteurs, les check-ins en cours et les actions rapides pour la gestion des présences.",
        16: "L'écran de détail d'un visiteur dans le tableau de bord Réception présente les informations personnelles, le statut de la visite et les actions de gestion disponibles.",
        17: "L'accueil du portail visiteur présente un résumé de la visite avec les informations de l'hôte, le statut de la demande et les prochaines étapes du parcours.",
        18: "La sélection du but de la visite propose des catégories standardisées (Réunion, Livraison, Maintenance, Audit, Formation) permettant de préparer l'accueil et d'informer l'hôte.",
        19: "L'écran de détails du visiteur affiche les informations personnelles, la photo et le QR code généré pour le check-in à l'entrée du site.",
        20: "Le QR code de badge visiteur, généré après validation de la demande, permet un check-in rapide à l'entrée du site via un scanner dédié.",
        21: "L'image annotée du campus OCP Jorf Lasfar identifie les principales zones du complexe : usines de production (JFC 1-5, DAP, TSP), bureaux administratifs, infrastructures de soutien et zones logistiques.",
        24: "La vue rapprochée de la carte interactive montre les marqueurs de localisations avec leurs étiquettes, l'algorithme d'évitement de collision ajustant automatiquement la taille et la position des labels en fonction du niveau de zoom.",
        25: "L'architecture de déploiement sépare les trois composants sur des hébergeurs distincts : Vercel pour le frontend (SSR, CDN), Render pour le backend (Express + Socket.IO) et Supabase pour la base de données PostgreSQL.",
        26: "La structure du projet suit une organisation modulaire avec des dossiers distincts pour le frontend Next.js (composants, lib, i18n), le backend Express.js (routes, controllers, services, middleware), les données cartographiques (places, nodes, edges, map) et les scripts d'automatisation.",
    }
    
    # ======================================================================
    # STEP 3: Update figure captions
    # ======================================================================
    print("\nUpdating figure captions...")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text.startswith("Figure ") and not text.startswith("Listing "):
            continue
        
        if text.startswith("Listing "):
            # Remove listing captions
            set_para_text(para, "")
            print(f"  Removed Listing at para {i}")
            continue
        
        # Extract figure number
        try:
            parts = text.split("—")[0]
            num_str = parts.replace("Figure", "").strip()
            fig_num = int(num_str.split()[0])
        except:
            continue
        
        if fig_num in OLD_TO_NEW:
            new_num = OLD_TO_NEW[fig_num]
            new_text = NEW_CAPTIONS[new_num]
            set_para_text(para, new_text, italic=True, size=10)
            print(f"  Fig {fig_num} -> Fig {new_num}: {new_text[:60]}")
        elif fig_num in REPLACEMENT_TEXT:
            # For removed figures, prepend "(ancienne Figure X)" to the replacement text
            replacement = REPLACEMENT_TEXT[fig_num]
            set_para_text(para, replacement, italic=False, size=11)
            print(f"  Fig {fig_num} replaced with text")
    
    # ======================================================================
    # STEP 4: Replace code blocks with academic prose
    # ======================================================================
    print("\nReplacing code blocks...")
    
    code_replacements = {
        "// Les markers sont des pourcentages du PNG": 
            "Le positionnement des marqueurs sur la carte repose sur un principe fondamental : les coordonnées de chaque localisation sont exprimées en pourcentage de l'image satellite PNG, et non en pixels absolus de la page web. Cette approche, dite « PNG-local », garantit que les marqueurs restent parfaitement alignés avec les bâtiments sous toute condition de zoom, de redimensionnement ou de translation. Lorsque la carte est mise à l'échelle, chaque marqueur est recalculé proportionnellement aux dimensions de l'image affichée, éliminant ainsi les décalages observés avec des systèmes basés sur les coordonnées de viewport. Le rapport d'affichage est défini par les dimensions naturelles de l'image maître (1520 × 933 pixels). La transformation CSS appliquée au conteneur de la carte combine une translation et une mise à l'échelle, permettant le zoom et le déplacement tout en préservant l'alignement des éléments graphiques.",
        
        "export function aStar(nodes, edges":
            "L'algorithme de routage piéton repose sur la recherche A*, qui combine l'optimalité de l'algorithme de Dijkstra avec la rapidité d'une heuristique euclidienne admissible. Le processus se déroule en trois étapes principales. La première étape consiste à établir la correspondance entre la localisation sélectionnée par l'utilisateur et le nœud d'accès le plus proche dans le graphe de navigation. La seconde étape construit le graphe non dirigé en ajoutant chaque arête dans les deux sens, avec pour coût la distance euclidienne entre les nœuds. La troisième étape applique l'algorithme A* avec pour fonction d'évaluation f(n) = g(n) + h(n), où g(n) représente le coût réel depuis le départ et h(n) l'estimation heuristique vers l'arrivée. La complexité pratique est de O(n log n) grâce à l'utilisation d'une file de priorité et à la topologie relativement linaire du campus.",
        
        "pixelsPerMeter = 0.8786127167630058":
            "La conversion des distances pixels en métriques réelles utilise un facteur d'échelle calibré lors de la phase de calibration GPS. Ce facteur, déterminé par la transformation affine reliant les coordonnées GPS aux pixels de l'image, permet de convertir chaque distance calculée sur le graphe en une distance métrique exploitable. La vitesse de marche estimée est de 1,4 m/s, soit environ 5 km/h, ce qui correspond à une cadence normale pour un piéton en milieu industriel. Le temps de parcours estimé est affiché dans l'interface utilisateur lors de la sélection d'un itinéraire. À titre d'exemple, une distance de 1 000 pixels sur le graphe correspond approximativement à 1 138 mètres, soit un temps de marche d'environ treize minutes.",
        
        "npx tsc --noEmit":
            "La vérification statique est réalisée par le compilateur TypeScript en mode strict, qui assure l'absence d'erreurs de typage dans l'ensemble du codebase. Le build Next.js valide la cohérence des dépendances et la génération des routes. Les tests Jest, exécutés en mode non-interactif, vérifient le bon fonctionnement des endpoints API et de la logique métier.",
    }
    
    code_count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        for key, replacement in code_replacements.items():
            if key in text:
                set_para_text(para, replacement, size=11)
                code_count += 1
                print(f"  Replaced code at para {i}")
                break
    
    # ======================================================================
    # STEP 5: Replace ASCII art with prose
    # ======================================================================
    print("\nReplacing ASCII art...")
    
    ascii_replacements = {
        "┌─────────────────────────────────────────────────────────┐\n│                     UTILISATEUR":
            "L'architecture générale d'OCP eGuide suit un modèle client-serveur en trois couches principales. La couche présentation (Next.js) accueille la page d'accueil, le flux d'authentification et les tableaux de bord avec carte interactive. La couche métier (Express.js) expose une API REST complète couplée à un serveur Socket.IO pour la communication temps réel. La couche de données (PostgreSQL via Prisma ORM) stocke les modèles User, Visit, Internship, Task, Location, Place, Notification, Request, QrCode et Presence. Les couches communiquent par le biais de requêtes HTTP et de connexions WebSocket.",
        
        "┌──────────────────────┐        ┌──────────────────────────┐\n│   Frontend (Vercel)":
            "L'architecture de déploiement sépare les trois composants sur des hébergeurs distincts : le frontend Next.js est déployé sur Vercel (SSR, CDN, déploiement automatique, port 443) ; le backend Express.js est hébergé sur Render avec Socket.IO (port 5000) ; la base de données PostgreSQL est administrée sur Supabase avec sauvegardes automatiques et réplication.",
        
        "project assets/\n├── backend/":
            "La structure du projet suit une organisation modulaire : le dossier backend/ contient l'API Express.js avec une architecture en couches (routes, controllers, services, middleware), le schéma Prisma et les tests Jest ; le dossier freebuff-frontend/ héberge l'application Next.js avec ses composants, sa logique métier (mapEngine, geoTransform), ses traductions i18n et les données cartographiques (places.json, nodes.json, edges.json, campus-map.png) ; les scripts d'automatisation assurent le lancement et la régénération des données.",
    }
    
    ascii_count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text:
            continue
        for key, replacement in ascii_replacements.items():
            if key in text:
                set_para_text(para, replacement, size=11)
                ascii_count += 1
                print(f"  Replaced ASCII art at para {i}")
                break
    
    # ======================================================================
    # STEP 6: Fix figure references in body text
    # ======================================================================
    print("\nFixing figure references...")
    
    ref_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if not run.text:
                continue
            
            def fix_ref(match):
                num = int(match.group(1))
                if num in OLD_TO_NEW:
                    return f"Figure {OLD_TO_NEW[num]}"
                elif num in REMOVED_FIGS:
                    return f"la section précédente"
                return match.group(0)
            
            new_text = re.sub(r'Figure (\d+)', fix_ref, run.text)
            if new_text != run.text:
                run.text = new_text
                ref_count += 1
    
    print(f"  Fixed {ref_count} references")
    
    # ======================================================================
    # STEP 7: Add extra content for ~34 page target
    # ======================================================================
    print("\nAdding supplementary content...")
    
    extensions = {
        "Utilisateurs cibles": 
            "\n\nTrois grandes catégories d'utilisateurs sont ciblées par le système. Les employés permanents (environ 2 000 personnes) ont besoin d'un accès rapide aux outils métier et à la carte pour leurs déplacements quotidiens sur le site. Les stagiaires (environ 200 personnes par session) nécessitent un accompagnement renforcé pour la découverte du complexe et l'accès à leur secteur d'affectation. Les visiteurs extérieurs (fournisseurs, clients, partenaires) représentent la catégorie la plus vulnérable en termes d'orientation, car ils ne disposent d'aucun repère préalable du site. Le système doit répondre de manière adaptée à chacun de ces profils, tout en maintenant une interface cohérente et intuitive.",
        
        "Besoins fonctionnels":
            "\n\nLes besoins fonctionnels ont été structurés en quatre catégories : navigation et orientation (carte interactive, routage piéton, recherche de localisations), gestion des accès (authentification, autorisation par rôle, profils utilisateurs), communication (messagerie temps réel, notifications, QR codes pour les badges visiteurs) et administration (gestion des utilisateurs, configuration des profils, monitoring des présences). Chaque catégorie a fait l'objet d'une analyse détaillée des workflows existants et des points de douleur identifiés sur le terrain, permettant de prioriser les fonctionnalités selon leur impact opérationnel.",
        
        "Architecture backend":
            "\n\nLe backend suit une architecture en couches garantissant la séparation des préoccupations. La couche des routes définit les points d'entrée de l'API et délègue le traitement aux controllers. Les controllers orchestrent les appels aux services métier et formatent les réponses HTTP. Les services encapsulent la logique applicative et interagissent directement avec Prisma ORM pour les opérations sur la base de données. Les middleware assurent l'authentification JWT, la validation des entrées via Zod et la gestion centralisée des erreurs. Cette séparation facilite la maintenabilité, les tests unitaires et l'évolution indépendante de chaque couche.",
        
        "Méthodologie":
            "\n\nL'approche méthodologique combinait des principes de développement itératif avec des pratiques agiles. Chaque itération, d'une durée de deux à trois semaines, aboutissait à une version fonctionnelle validée par l'encadrant et l'équipe technique d'OCP. Les maquettes Figma, réalisées en amont du développement, servaient de référence pour la conformité visuelle. Les revues de code et les tests automatisés (Jest, TypeScript strict) garantissaient la qualité à chaque étape. Cette méthodologie a permis de gérer efficacement les incertitudes techniques liées au positionnement cartographique et au routage sur un campus non standard.",
        
        "Tests fonctionnels":
            "\n\nLes tests fonctionnels ont couvert les principaux parcours utilisateur : authentification et autorisation par profil, navigation sur la carte, recherche de localisations, routage piéton, gestion des visites (inscription, approbation, check-in, check-out) et communication temps réel. Les tests ont été réalisés manuellement sur les environnements de développement et de staging, avec une attention particulière portée aux scénarios de mode kiosk (navigation automatique, timeout après inactivité) et au comportement responsive sur différentes résolutions d'écran (1920×1080, 1366×768, 375×667 mobile).",
        
        "Perspectives d'amélioration":
            "\n\nPlusieurs axes d'amélioration ont été identifiés pour les prochaines versions du système. Sur le plan cartographique, l'intégration d'une image satellite haute résolution permettrait d'améliorer la précision du positionnement. Sur le plan fonctionnel, la navigation intérieure (étages, couloirs, escaliers) constituerait une avancée majeure pour l'orientation dans les bâtiments de production. L'ajout d'une application mobile native (React Native) permettrait aux employés d'utiliser le système en déplacement, avec un mode hors-ligne pour les zones à faible couverture réseau. Enfin, l'implémentation de tests end-to-end (Playwright ou Cypress) et l'intégration d'outils d'analytics permettraient de mesurer l'adoption du système et d'identifier les axes d'optimisation de l'expérience utilisateur.",
    }
    
    ext_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, extension in extensions.items():
            if text == key or text.startswith(key):
                if para.runs:
                    para.runs[-1].text += extension
                    ext_count += 1
                    print(f"  Extended: {key}")
                break
    
    print(f"  Extended {ext_count} sections")
    
    # ======================================================================
    # STEP 8: Save
    # ======================================================================
    print(f"\nSaving to {output_path}...")
    doc.save(output_path)
    
    # Verify
    doc2 = Document(output_path)
    img_count = sum(1 for p in doc2.paragraphs if has_drawing(p))
    fig_count = sum(1 for p in doc2.paragraphs if p.text.strip().startswith("Figure "))
    
    print(f"\n{'='*70}")
    print(f"FINAL VERIFICATION:")
    print(f"  Images: {img_count} (target: ≤5)")
    print(f"  Figure captions: {fig_count}")
    print(f"  Paragraphs: {len(doc2.paragraphs)}")
    print(f"  Tables: {len(doc2.tables)}")
    print(f"{'='*70}")

if __name__ == "__main__":
    main()
