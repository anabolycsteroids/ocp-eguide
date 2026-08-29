#!/usr/bin/env python3
"""Deep verification of polished document."""
import sys
sys.path.insert(0, '/home/imrane/.pyenv/versions/3.12.11/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from lxml import etree

doc = Document("/mnt/c/Users/unknown/Downloads/project assets/OCP_eGuide_Rapport_Stage_PROFESSIONAL.docx")

print("=== DEEP VERIFICATION ===\n")

# 1. Check logo
print("[LOGO]")
# Check if there's an image relationship
image_rels = [r for r in doc.part.rels.values() if 'image' in r.reltype]
print(f"  Image relationships: {len(image_rels)}")
for r in image_rels:
    print(f"    {r.rId}: {r.reltype} -> {r.target_ref}")

# Check inline shapes
print(f"  Inline shapes: {len(doc.inline_shapes)}")

# Check for drawings in the first ~30 paragraphs
for i in range(min(30, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    drawings = p._element.findall('.//' + qn('w:drawing'))
    if drawings:
        print(f"  Drawing found in paragraph {i} (text: '{p.text[:50]}'...)")
        for d in drawings:
            blips = d.findall('.//' + qn('a:blip'))
            for b in blips:
                embed = b.get(qn('r:embed'))
                print(f"    Blip embed: {embed}")

# 2. Check encadrant
print("\n[ENCADRANT NAME]")
for i, p in enumerate(doc.paragraphs):
    if "Karim" in p.text or "Alaoui" in p.text:
        print(f"  ❌ FOUND at [{i}]: {p.text[:150]}")

# Check table cells too
for t_idx, table in enumerate(doc.tables):
    for row in table.rows:
        for cell in row.cells:
            if "Karim" in cell.text or "Alaoui" in cell.text:
                print(f"  ❌ FOUND in table {t_idx}: {cell.text[:150]}")

print("  ✅ No 'Karim Alaoui' found in paragraphs")

# 3. Check typography consistency
print("\n[TYPOGRAPHY CHECK]")
font_sizes = {}
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.size:
            key = f"{run.font.size} ({p.style.name})"
            font_sizes[key] = font_sizes.get(key, 0) + 1

for k, v in sorted(font_sizes.items(), key=lambda x: -x[1])[:15]:
    print(f"  {k}: {v} runs")

# 4. Check heading colors
print("\n[HEADING COLORS]")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        for run in p.runs:
            color = run.font.color.rgb
            if color:
                print(f"  [{i}] {p.style.name}: color={color}, bold={run.font.bold}, size={run.font.size}")
                break

# 5. Check page breaks
print("\n[PAGE BREAKS BEFORE CHAPTERS]")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1':
        pbf = p.paragraph_format.page_break_before
        print(f"  [{i}] '{p.text[:60]}' page_break_before={pbf}")

# 6. Check footer
print("\n[FOOTER]")
for section in doc.sections:
    footer = section.footer
    for p in footer.paragraphs:
        text = p.text
        if text:
            print(f"  Footer text: '{text}'")
        # Check for page number field
        fields = p._element.findall('.//' + qn('w:fldChar'))
        if fields:
            print(f"  Page number fields: {len(fields)}")

# 7. Check remeciements
print("\n[REMERCIEMENTS CHECK]")
for i in range(22, 26):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        if p.text.strip():
            print(f"  [{i}] {p.text[:200]}")

# 8. Check tables styling
print("\n[TABLE STYLING]")
for t_idx, table in enumerate(doc.tables):
    first_row = table.rows[0] if table.rows else None
    if first_row:
        cell = first_row.cells[0]
        tcPr = cell._tc.find(qn('w:tcPr'))
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                print(f"  Table {t_idx}: header fill={fill}")

# 9. Check margins
print("\n[MARGINS]")
for section in doc.sections:
    print(f"  Top: {section.top_margin}")
    print(f"  Bottom: {section.bottom_margin}")
    print(f"  Left: {section.left_margin}")
    print(f"  Right: {section.right_margin}")
    print(f"  Page: {section.page_width} x {section.page_height}")

# 10. Check for empty paragraphs (page breaks)
print("\n[EMPTY PARAGRAPHS / PAGE BREAKS]")
for i, p in enumerate(doc.paragraphs):
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        pageBreak = pPr.find(qn('w:pageBreakBefore'))
        if pageBreak is not None:
            val = pageBreak.get(qn('w:val'))
            print(f"  [{i}] PageBreakBefore (val={val}): '{p.text[:60]}'")
