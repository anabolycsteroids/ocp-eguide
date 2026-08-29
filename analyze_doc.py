#!/usr/bin/env python3
import sys
sys.path.insert(0, '/home/imrane/.pyenv/versions/3.12.11/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Inches, Pt, Cm
from collections import Counter

doc = Document("/mnt/c/Users/unknown/Downloads/project assets/OCP_eGuide_Rapport_Stage_Final_Clean.docx")

print("=== DOCUMENT ANALYSIS ===")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")
print(f"Total sections: {len(doc.sections)}")

# Analyze styles
style_counts = Counter()
for p in doc.paragraphs:
    style_counts[p.style.name] += 1

print("\n=== STYLE USAGE ===")
for style, count in style_counts.most_common(30):
    print(f"  {style}: {count}")

# Show first 100 paragraphs with their styles
print("\n=== FIRST 100 PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs[:100]):
    text = p.text.strip()
    if text:
        print(f"  [{i}] ({p.style.name}) {text[:120]}")
    elif p.style.name.startswith('Heading'):
        print(f"  [{i}] ({p.style.name}) [empty heading]")

# Show images
print(f"\n=== IMAGES: {len(doc.inline_shapes)} ===")
for i, shape in enumerate(doc.inline_shapes):
    print(f"  Image {i}: type={shape.type}, width={shape.width}, height={shape.height}")

# Show headers/footers
print("\n=== SECTIONS ===")
for i, section in enumerate(doc.sections):
    print(f"  Section {i}: width={section.page_width}, height={section.page_height}")
    print(f"    margins: top={section.top_margin}, bottom={section.bottom_margin}, left={section.left_margin}, right={section.right_margin}")
    if section.header:
        header_text = " ".join([p.text for p in section.header.paragraphs]).strip()
        if header_text:
            print(f"    Header: {header_text[:100]}")
    if section.footer:
        footer_text = " ".join([p.text for p in section.footer.paragraphs]).strip()
        if footer_text:
            print(f"    Footer: {footer_text[:100]}")

# Search for encadrant name
print("\n=== SEARCHING FOR 'Karim Alaoui' ===")
for i, p in enumerate(doc.paragraphs):
    if "Karim" in p.text or "Alaoui" in p.text:
        print(f"  [{i}] ({p.style.name}) {p.text[:200]}")

print("\n=== SEARCHING FOR 'encadrant' (case-insensitive) ===")
for i, p in enumerate(doc.paragraphs):
    if "encadrant" in p.text.lower() or "Encadrant" in p.text:
        print(f"  [{i}] ({p.style.name}) {p.text[:200]}")

# Show all headings
print("\n=== ALL HEADINGS ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading'):
        print(f"  [{i}] ({p.style.name}) {p.text[:150]}")

# Show tables
print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    print(f"  Table {t_idx}: {rows} rows x {cols} cols")
    if rows > 0:
        header_cells = [cell.text.strip()[:30] for cell in table.rows[0].cells]
        print(f"    Headers: {header_cells}")
