#!/usr/bin/env python3
"""
OCP eGuide Report Polisher
Transforms the Word document into a professional final-year engineering report.
"""
import sys
sys.path.insert(0, '/home/imrane/.pyenv/versions/3.12.11/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from lxml import etree
import os
import copy

# Paths
INPUT_PATH = "/mnt/c/Users/unknown/Downloads/project assets/OCP_eGuide_Rapport_Stage_Final_Clean.docx"
OUTPUT_PATH = "/mnt/c/Users/unknown/Downloads/project assets/OCP_eGuide_Rapport_Stage_PROFESSIONAL.docx"
LOGO_PATH = "/mnt/c/Users/unknown/Downloads/project assets/ocp_logo_for_doc.png"

# OCP Brand Colors
OCP_GREEN = RGBColor(0x00, 0x82, 0x3B)      # Primary OCP green
OCP_DARK = RGBColor(0x1B, 0x2A, 0x4A)       # Dark navy
OCP_ACCENT = RGBColor(0x00, 0x6B, 0x33)     # Darker green
OCP_LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2) # Light background
OCP_MED_GRAY = RGBColor(0x66, 0x66, 0x66)   # Medium gray for captions
OCP_TABLE_HEADER_BG = "00823B"                # Green table headers
OCP_TABLE_ALT_BG = "F0F7F2"                  # Light green alternating rows
OCP_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
OCP_BLACK = RGBColor(0x00, 0x00, 0x00)

# Font settings
BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"
BODY_SIZE = Pt(11)
LINE_SPACING = 1.15

print("=" * 60)
print("OCP eGuide Report Polisher")
print("=" * 60)

# Load document
doc = Document(INPUT_PATH)
print(f"Loaded document with {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# ============================================================
# SECTION 1: Fix document-level settings and margins
# ============================================================
print("\n[1/12] Fixing page layout and margins...")

for section in doc.sections:
    section.page_width = Cm(21)      # A4 width
    section.page_height = Cm(29.7)   # A4 height
    section.top_margin = Cm(2.5)     # Professional margins
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# SECTION 2: Setup heading styles with OCP colors
# ============================================================
print("[2/12] Setting up professional typography...")

def setup_paragraph_spacing(paragraph, before=0, after=6, line_spacing=1.15):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing

def set_run_font(run, font_name=BODY_FONT, size=BODY_SIZE, bold=False, italic=False, color=None):
    """Configure a run's font."""
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    # Ensure East Asian font fallback
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)

def style_heading(paragraph, level=1):
    """Apply consistent heading style."""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.2
    
    for run in paragraph.runs:
        if level == 1:
            run.font.name = HEADING_FONT
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = OCP_DARK
        elif level == 2:
            run.font.name = HEADING_FONT
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = OCP_GREEN
        elif level == 3:
            run.font.name = HEADING_FONT
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = OCP_ACCENT

# Apply styles to all paragraphs
for i, para in enumerate(doc.paragraphs):
    style_name = para.style.name
    
    if style_name == 'Heading 1':
        style_heading(para, 1)
        setup_paragraph_spacing(para, before=24, after=12, line_spacing=1.2)
        
    elif style_name == 'Heading 2':
        style_heading(para, 2)
        setup_paragraph_spacing(para, before=16, after=8, line_spacing=1.2)
        
    elif style_name == 'Heading 3':
        style_heading(para, 3)
        setup_paragraph_spacing(para, before=12, after=6, line_spacing=1.2)
        
    elif style_name == 'Normal':
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        setup_paragraph_spacing(para, before=0, after=6, line_spacing=1.15)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    elif style_name == 'List Bullet':
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = BODY_SIZE
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        setup_paragraph_spacing(para, before=2, after=2, line_spacing=1.15)

# ============================================================
# SECTION 3: Remove Karim Alaoui name
# ============================================================
print("[3/12] Removing unconfirmed supervisor name...")

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    # Fix cover page encadrant
    if "Karim Alaoui" in text:
        print(f"  Found 'Karim Alaoui' at paragraph {i}")
        for run in para.runs:
            if "Karim Alaoui" in run.text:
                run.text = run.text.replace("M. Karim Alaoui", "______________________")
                print(f"    → Replaced in run")
    
    # Fix remerciements
    if "mon encadrant, M. Karim" in para.text:
        for run in para.runs:
            if "M. Karim" in run.text:
                run.text = run.text.replace("M. Karim Alaoui", "")
                # Clean up extra spaces/comma
        # Need to also fix the surrounding text
        full_text = para.text
        if "mon encadrant," in full_text and "pour sa disponibilité" not in full_text:
            # Replace the whole sentence about encadrant
            for run in para.runs:
                if "mon encadrant" in run.text:
                    run.text = run.text.replace("mon encadrant, ", "mon encadrant ")
    
    # Fix chapter 1 supervision reference
    if "sous la supervision de M. Karim Alaoui" in para.text:
        for run in para.runs:
            if "M. Karim Alaoui" in run.text:
                run.text = run.text.replace("sous la supervision de M. Karim Alaoui.", "sous la supervision de mon encadrant.")
                print(f"    → Fixed supervision reference")

# ============================================================
# SECTION 4: Fix typographical errors
# ============================================================
print("[4/12] Fixing typographical errors...")

typo_fixes = [
    ("Socket.IOpour", "Socket.IO pour"),
    ("Node. js", "Node.js"),
    ("place. json", "places.json"),
    ("nodes. json", "nodes.json"),
    ("edges. json", "edges.json"),
    ("campus-map. png", "campus-map.png"),
    ("generate-map-data. js", "generate-map-data.js"),
    ("regenerate-map. mjs", "regenerate-map.mjs"),
    ("convert-coordinates. js", "convert-coordinates.js"),
    ("Schema Prisma", "Schéma Prisma"),
    ("indéfectiblee", "indéfectible"),
    ("permettrent", "permettent"),
    ("Node .js", "Node.js"),
    ("TypeScript .", "TypeScript."),
    ("Prisma.", "Prisma."),
    ("Express .js", "Express.js"),
    ("Tailwind .CSS", "Tailwind CSS"),
    ("React .", "React."),
]

for para in doc.paragraphs:
    for run in para.runs:
        for old, new in typo_fixes:
            if old in run.text:
                run.text = run.text.replace(old, new)
                print(f"  Fixed: '{old}' → '{new}'")

# Also fix in tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for old, new in typo_fixes:
                        if old in run.text:
                            run.text = run.text.replace(old, new)

# ============================================================
# SECTION 5: Fix cover page
# ============================================================
print("[5/12] Polishing cover page...")

# The cover page is paragraphs 0-20 approximately
# We need to rebuild the cover page area

# First, let's identify the cover page area
cover_start = 0
cover_end = 21  # Before Remerciements

# Clear and rebuild paragraphs 6-20 (the actual cover content)
# We'll work with what's there and restyle it

for i in range(cover_start, cover_end):
    para = doc.paragraphs[i]
    text = para.text.strip()
    
    if i == 6:  # "OCP eGuide"
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(4)
        for run in para.runs:
            run.font.name = HEADING_FONT
            run.font.size = Pt(28)
            run.font.bold = True
            run.font.color.rgb = OCP_DARK
            
    elif i == 7:  # "Système Numérique Interactif..."
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(20)
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(13)
            run.font.italic = True
            run.font.color.rgb = OCP_MED_GRAY
            
    elif i == 10:  # "Rapport de Stage"
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(20)
        for run in para.runs:
            run.font.name = HEADING_FONT
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = OCP_GREEN
            
    elif i >= 13 and i <= 18:  # Cover details
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
    elif i == 18:  # Année universitaire
        para.paragraph_format.space_before = Pt(10)
        
    elif i == 20:  # "Août 2026"
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after = Pt(0)
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(12)
            run.font.color.rgb = OCP_MED_GRAY

# Add logo to cover page
# Insert the logo as a picture at the beginning
# We'll add it as the first element in paragraph 6's run area
# Actually, let's insert it before paragraph 6

# Access the XML to insert logo
body = doc.element.body
cover_para_6 = doc.paragraphs[6]._element

# Create a new paragraph for the logo
logo_para = parse_xml(
    f'<w:p {nsdecls("w")}>'
    f'  <w:pPr>'
    f'    <w:jc w:val="center"/>'
    f'    <w:spacing w:before="0" w:after="200"/>'
    f'  </w:pPr>'
    f'  <w:r>'
    f'    <w:drawing>'
    f'      <wp:inline xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    f'        distT="0" distB="0" distL="0" distR="0">'
    f'        <wp:extent cx="2400000" cy="2400000"/>'  # ~6.35cm x 6.35cm
    f'        <wp:docPr id="1" name="OCP Logo"/>'
    f'        <a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
    f'          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
    f'            <pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
    f'              <pic:nvPicPr>'
    f'                <pic:cNvPr id="1" name="OCP Logo"/>'
    f'                <pic:cNvPicPr/>'
    f'              </pic:nvPicPr>'
    f'              <pic:blipFill>'
    f'                <a:blip r:embed="rId9999" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
    f'              </pic:blipFill>'
    f'              <pic:spPr>'
    f'                <a:xfrm>'
    f'                  <a:off x="0" y="0"/>'
    f'                  <a:ext cx="2400000" cy="2400000"/>'
    f'                </a:xfrm>'
    f'                <a:prstGeom prst="rect"/>'
    f'              </pic:spPr>'
    f'            </pic:pic>'
    f'          </a:graphicData>'
    f'        </a:graphic>'
    f'      </wp:inline>'
    f'    </w:drawing>'
    f'  </w:r>'
    f'</w:p>'
)

# Insert logo before paragraph 6
cover_para_6.addprevious(logo_para)

# Now we need to add the relationship for the logo image
# We'll do this by adding the image part to the document
import docx.opc.constants
from docx.opc.part import Part
from docx.opc.packuri import PackURI

# Add image as a relationship
doc_part = doc.part
logo_part_name = PackURI('/word/media/ocp_logo.png')

# Read the PNG file
with open(LOGO_PATH, 'rb') as f:
    logo_data = f.read()

# Add the image part
from docx.image.image import Image
from docx.opc.part import Part
import docx.opc.constants as constants

# Create the image part properly
logo_part = Part(
    logo_part_name,
    'image/png',
    logo_data,
    doc_part.package
)

# Add relationship
r_id = doc_part.relate_to(logo_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image')

# Update the drawing to use the correct rId
drawing = logo_para.find(qn('w:r')).find(qn('w:drawing'))
blip = drawing.find('.//' + qn('a:blip'))
if blip is not None:
    blip.set(qn('r:embed'), r_id)

print(f"  Logo added with relationship ID: {r_id}")

# ============================================================
# SECTION 6: Fix remerciements
# ============================================================
print("[6/12] Fixing remerciements...")

# Find and fix the remerciements paragraph
for i, para in enumerate(doc.paragraphs):
    if i == 23:  # The remerciements body
        text = para.text
        # Replace the sentence about encadrant
        for run in para.runs:
            if "mon encadrant," in run.text and "Karim" not in run.text:
                # Already cleaned, but let's make it generic
                pass
            if "mon encadrant" in run.text:
                # Make it generic - remove specific name references
                if "pour sa disponibilité" not in run.text:
                    # Replace the full sentence
                    old_text = run.text
                    # Generic replacement
                    new_text = old_text.replace(
                        "Je remercie tout d'abord mon encadrant, ",
                        "Je remercie tout d'abord mon encadrant "
                    )
                    if new_text != old_text:
                        run.text = new_text

# ============================================================
# SECTION 7: Professionalize tables
# ============================================================
print("[7/12] Professionalizing tables...")

for t_idx, table in enumerate(doc.tables):
    # Set table alignment
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style the table
    tbl = table._tbl
    
    # Remove default borders and add clean ones
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        tbl.insert(0, tblPr)
    
    # Set table borders
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
      <w:top w:val="single" w:sz="4" w:space="0" w:color="00823B"/>
      <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:bottom w:val="single" w:sz="4" w:space="0" w:color="00823B"/>
      <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
      <w:insideH w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>
      <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
    </w:tblBorders>
    '''
    # Remove existing borders
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    tblPr.append(parse_xml(borders_xml))
    
    # Style each row
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            # Style cell paragraphs
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.line_spacing = 1.0
                
                for run in para.runs:
                    run.font.name = BODY_FONT
                    run.font.size = Pt(9.5)
                    
                    if r_idx == 0:  # Header row
                        run.font.bold = True
                        run.font.color.rgb = OCP_WHITE
                        run.font.size = Pt(10)
                    else:
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            
            # Header row background
            if r_idx == 0:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{OCP_TABLE_HEADER_BG}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            elif r_idx % 2 == 0:  # Alternating rows
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{OCP_TABLE_ALT_BG}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)

# ============================================================
# SECTION 8: Add page breaks before major chapters
# ============================================================
print("[8/12] Adding page breaks before chapters...")

# Find all Heading 1 paragraphs and ensure page break before
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'Heading 1':
        pf = para.paragraph_format
        pf.page_break_before = True

# ============================================================
# SECTION 9: Style figure captions
# ============================================================
print("[9/12] Fixing figure captions...")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("Figure ") and "—" in text:
        # This is a figure caption
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = OCP_MED_GRAY
    elif text.startswith("Figure ") and ("-" in text or "–" in text):
        # Also catch figure captions with regular dash
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        for run in para.runs:
            run.font.name = BODY_FONT
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = OCP_MED_GRAY

# ============================================================
# SECTION 10: Style special paragraphs (keywords, code references)
# ============================================================
print("[10/12] Styling special elements...")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Keywords line
    if text.startswith("Mots-clés"):
        for run in para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = OCP_MED_GRAY
    
    # "Table des matières" placeholder
    if i == 34:  # The TOC placeholder text
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = OCP_MED_GRAY

# ============================================================
# SECTION 11: Headers and Footers
# ============================================================
print("[11/12] Adding professional headers and footers...")

for section in doc.sections:
    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    
    # Clear existing footer
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""
    
    if footer.paragraphs:
        footer_para = footer.paragraphs[0]
    else:
        footer_para = footer.add_paragraph()
    
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add footer text
    run = footer_para.add_run("OCP eGuide  •  Génie Informatique  •  2025–2026")
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = OCP_MED_GRAY
    
    # Add tab and page number
    footer_para.add_run("        ")
    run_page = footer_para.add_run()
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = OCP_MED_GRAY
    
    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_page._r.append(fldChar1)
    
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_page._r.append(instrText)
    
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_page._r.append(fldChar2)
    
    # Add a subtle line above footer
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="00823B"/>'
        f'</w:pBdr>'
    )
    footer_para._p.get_or_add_pPr().append(pBdr)

# ============================================================
# SECTION 12: Final document metadata and cleanup
# ============================================================
print("[12/12] Final cleanup and metadata...")

# Set document properties
core_props = doc.core_properties
core_props.author = "Imrane Belkoufa"
core_props.title = "OCP eGuide — Rapport de Stage"
core_props.subject = "Système Numérique Interactif d'Orientation du Campus"
core_props.keywords = "OCP eGuide, navigation campus, carte interactive, pathfinding, Next.js, Express.js"

# ============================================================
# SAVE
# ============================================================
print("\n" + "=" * 60)
print("Saving polished document...")
doc.save(OUTPUT_PATH)
print(f"✅ Saved to: {OUTPUT_PATH}")

# Verify
doc2 = Document(OUTPUT_PATH)
print(f"Verification: {len(doc2.paragraphs)} paragraphs, {len(doc2.tables)} tables")
print(f"Sections: {len(doc2.sections)}")

# Check encadrant name is gone
found_name = False
for p in doc2.paragraphs:
    if "Karim Alaoui" in p.text:
        found_name = True
        print(f"  ⚠️ WARNING: 'Karim Alaoui' still found in paragraph: {p.text[:100]}")

for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            if "Karim Alaoui" in cell.text:
                found_name = True
                print(f"  ⚠️ WARNING: 'Karim Alaoui' still found in table: {cell.text[:100]}")

if not found_name:
    print("  ✅ No 'Karim Alaoui' references found - name successfully removed")

print("\n✅ Document polishing complete!")
