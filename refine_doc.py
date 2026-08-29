#!/usr/bin/env python3
"""Final refinement pass - fix remaining issues."""
import sys
sys.path.insert(0, '/home/imrane/.pyenv/versions/3.12.11/lib/python3.12/site-packages')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

INPUT_PATH = "/mnt/c/Users/unknown/Downloads/project assets/OCP_eGuide_Rapport_Stage_PROFESSIONAL.docx"
OUTPUT_PATH = "/mnt/c/Users/unknown/Downloads/project assets/OCP_eGuide_Rapport_Stage_PROFESSIONAL.docx"

OCP_GREEN = RGBColor(0x00, 0x82, 0x3B)
OCP_DARK = RGBColor(0x1B, 0x2A, 0x4A)
OCP_TABLE_HEADER_BG = "00823B"
OCP_TABLE_ALT_BG = "F0F7F2"

doc = Document(INPUT_PATH)
print("Loaded for refinement pass...")

# ============================================================
# FIX 1: Table header colors - ensure correct OCP green
# ============================================================
print("\n[FIX 1] Ensuring correct table header colors...")
for t_idx, table in enumerate(doc.tables):
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # Remove any existing shading
            existing_shd = tcPr.find(qn('w:shd'))
            if existing_shd is not None:
                tcPr.remove(existing_shd)
            
            if r_idx == 0:  # Header row
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{OCP_TABLE_HEADER_BG}" w:val="clear"/>'
                )
                tcPr.append(shd)
                # Set header text to white
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
            elif r_idx % 2 == 0:  # Alternating rows
                shd = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{OCP_TABLE_ALT_BG}" w:val="clear"/>'
                )
                tcPr.append(shd)

# ============================================================
# FIX 2: Fix "Table des matières" heading color
# ============================================================
print("[FIX 2] Fixing Table des matières heading...")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "Table des matières":
        for run in para.runs:
            run.font.color.rgb = OCP_DARK
            print(f"  Fixed heading color at paragraph {i}")

# ============================================================
# FIX 3: Fix cover page "Rapport de Stage" paragraph
# ============================================================
print("[FIX 3] Checking cover page details...")
# Check that cover page paragraphs are properly styled
for i in range(20):
    para = doc.paragraphs[i]
    text = para.text.strip()
    if text:
        print(f"  [{i}] '{text[:60]}' - align={para.paragraph_format.alignment}")

# ============================================================
# FIX 4: Ensure all list bullets are properly styled
# ============================================================
print("[FIX 4] Styling list bullets...")
for i, para in enumerate(doc.paragraphs):
    if para.style.name == 'List Bullet':
        for run in para.runs:
            if not run.font.name:
                run.font.name = "Calibri"
            if not run.font.size:
                run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)

# ============================================================
# FIX 5: Add proper spacing between cover elements
# ============================================================
print("[FIX 5] Cover page spacing...")
# Ensure paragraph 5 (empty before logo) is clean
para5 = doc.paragraphs[5]
if not para5.text.strip():
    # This is likely an empty paragraph before cover content
    pf = para5.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

# ============================================================
# FIX 6: Style the "Résumé" keywords line
# ============================================================
print("[FIX 6] Styling résumé keywords...")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("Mots-clés"):
        for run in para.runs:
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        print(f"  Styled keywords at paragraph {i}")

# ============================================================
# FIX 7: Ensure consistent paragraph alignment
# ============================================================
print("[FIX 7] Checking paragraph alignment...")
centered_paragraphs = []
for i, para in enumerate(doc.paragraphs):
    if para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        text = para.text.strip()
        if text:
            centered_paragraphs.append((i, text[:60]))

print(f"  Found {len(centered_paragraphs)} centered paragraphs:")
for idx, text in centered_paragraphs[:10]:
    print(f"    [{idx}] '{text}'")

# ============================================================
# FIX 8: Fix any remaining spacing issues
# ============================================================
print("[FIX 8] Final spacing adjustments...")
# Ensure no paragraph has excessive spacing
for i, para in enumerate(doc.paragraphs):
    pf = para.paragraph_format
    if pf.space_before and pf.space_before > Pt(30):
        pf.space_before = Pt(24)
    if pf.space_after and pf.space_after > Pt(18):
        pf.space_after = Pt(12)

# ============================================================
# SAVE
# ============================================================
print("\n" + "=" * 60)
print("Saving refined document...")
doc.save(OUTPUT_PATH)
print(f"✅ Saved to: {OUTPUT_PATH}")

# Final verification
doc2 = Document(OUTPUT_PATH)
print(f"\nFinal verification:")
print(f"  Paragraphs: {len(doc2.paragraphs)}")
print(f"  Tables: {len(doc2.tables)}")

# Check encadrant one more time
found = False
for p in doc2.paragraphs:
    if "Karim" in p.text:
        found = True
        print(f"  ❌ Still found: {p.text[:100]}")
if not found:
    print("  ✅ Encadrant name fully removed")

# Check logo
image_count = len(doc2.inline_shapes)
print(f"  Images: {image_count} (includes logo + original 5)")

# Check footer
for section in doc2.sections:
    footer_text = "".join([p.text for p in section.footer.paragraphs])
    print(f"  Footer: '{footer_text[:80]}'")

print("\n✅ Refinement complete!")
